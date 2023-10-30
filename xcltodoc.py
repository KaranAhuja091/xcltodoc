import streamlit as st
import openpyxl
import requests
from bs4 import BeautifulSoup
from docx import Document
import os

# Function to process the Excel file
def process_excel(excel_file):
    workbook = openpyxl.load_workbook(excel_file)
    doc = Document()

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]

        for row in sheet.iter_rows(min_col=2, max_col=3, values_only=True):
            content = row[0]
            url = row[1]

            if content and url:
                if url is not None and url.strip() != "":
                    response = requests.get(url)

                    if response.status_code == 200:
                        soup = BeautifulSoup(response.text, 'html.parser')
                        title = soup.find('h1')

                        if content:
                            doc.add_paragraph(content)

                        if title:
                            doc.add_heading(title.text, level=1)

                        paragraphs = soup.find_all('p')
                        for paragraph in paragraphs:
                            doc.add_paragraph(paragraph.get_text())

                        if url:
                            doc.add_paragraph(f"Source URL: {url}")

    output_file_path = "output_document.docx"
    doc.save(output_file_path)

# Streamlit app
def main():
    st.title("Excel to Word Processor")

    uploaded_file = st.file_uploader("Upload an Excel file", type=["xlsx"])

    if uploaded_file is not None:
        st.write("File Uploaded: ", uploaded_file.name)

        if st.button("Process and Download"):
            process_excel(uploaded_file)
            st.success("Processing complete. Click below to download the Word document.")

            # Provide a download link to the Word document
            with open("output_document.docx", "rb") as doc_file:
                st.download_button("Download Word Document", doc_file)

if __name__ == "__main__":
    main()
